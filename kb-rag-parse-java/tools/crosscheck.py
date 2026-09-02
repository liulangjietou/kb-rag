"""两实现对拍：同一份样例字节分别发给 Python 的 kb-rag-parser 与本仓库的
kb-rag-parse-java，逐项比对契约字段是否一致。

这是「两套实现行为等价」这一说法的证据来源。单元测试只能证明各自符合自己的预期，
对拍才能证明两者对同一份输入给出同一个答案——尤其是 pdf 文本层、markdown 表格渲染、
时间戳归一这些"看起来显然、实际很容易悄悄漂移"的地方。

样例全部由本脚本用 Python 侧的库现场生成，两个服务吃的是同一份字节。

比对口径：
- /api/v1/parse：code、pages 数量与 page_no/scanned/ocr_source、按空白归一后的 text、
  images 的数量与 image_id/page_no/kind、markdown 中占位符个数、warnings 条数，
  以及两边各自的"逐页 markdown 拼接 == 合并 markdown"不变量。
  图片二进制不逐字节比对：两侧的图片编码链路本就不同（PyMuPDF 直出原始编码流，
  PDFBox 对非 JPEG 统一重编码为 PNG），契约要求的是"同一张图被报告一次、位置正确"，
  不是"字节相同"。
- /api/v1/parse/chat：skipped 计数与 sessions/messages 逐字段全等（含 msg_id 与 send_time）。

用法：

    # 终端 1：Python 实现
    cd kb-rag-parser && .venv/bin/uvicorn app.main:app --port 20012

    # 终端 2：Java 实现
    cd kb-rag-parse-java && java -jar target/kb-rag-parse-java-*.jar --server.port=20011

    # 终端 3
    kb-rag-parser/.venv/bin/python kb-rag-parse-java/tools/crosscheck.py

依赖 Python 侧的 pymupdf / python-docx / openpyxl 来生成样例，所以用
kb-rag-parser 的 .venv 跑最省事。退出码非 0 表示存在差异。

Author: owlzhangfq@gmail.com
"""
import io, csv, json, os, sys, urllib.request, uuid

PY_URL = os.getenv("PY_PARSER_URL", "http://127.0.0.1:20012")
JAVA_URL = os.getenv("JAVA_PARSER_URL", "http://127.0.0.1:20011")

def post_multipart(base, path, filename, content, fields):
    boundary = uuid.uuid4().hex
    body = io.BytesIO()
    def w(s): body.write(s.encode('utf-8') if isinstance(s, str) else s)
    for k, v in fields.items():
        if v is None: continue
        w(f"--{boundary}\r\nContent-Disposition: form-data; name=\"{k}\"\r\n\r\n{v}\r\n")
    w(f"--{boundary}\r\nContent-Disposition: form-data; name=\"file\"; filename=\"{filename}\"\r\n")
    w("Content-Type: application/octet-stream\r\n\r\n")
    w(content); w(f"\r\n--{boundary}--\r\n")
    req = urllib.request.Request(base + path, data=body.getvalue(),
        headers={"Content-Type": f"multipart/form-data; boundary={boundary}"})
    with urllib.request.urlopen(req, timeout=120) as resp:
        return json.loads(resp.read().decode('utf-8'))

def parse_doc(base, filename, content, ext):
    return post_multipart(base, "/api/v1/parse", filename, content, {"file_ext": ext})

def parse_chat(base, filename, content, ext, profile=None, yml=None):
    return post_multipart(base, "/api/v1/parse/chat", filename, content,
                          {"file_ext": ext, "mapping_profile": profile, "profile_yaml": yml})

# ---------- 样例生成（用 Python 侧的库，两边吃同一份字节） ----------
import pymupdf
from docx import Document
from openpyxl import Workbook

def pdf_bytes(text="Hello kb-rag PDF, this page has a normal text layer."):
    d = pymupdf.open(); p = d.new_page(); p.insert_text((72,72), text)
    b = d.tobytes(); d.close(); return b

def multi_pdf(n=3):
    d = pymupdf.open()
    for i in range(1, n+1):
        p = d.new_page(); p.insert_text((72,72), f"Page {i} body text, long enough to have a real text layer.")
    b = d.tobytes(); d.close(); return b

def tiny_png():
    d = pymupdf.open(); p = d.new_page(width=40, height=40)
    b = p.get_pixmap().tobytes("png"); d.close(); return b

def pdf_with_image():
    d = pymupdf.open(); p = d.new_page()
    p.insert_text((72,72), "Hello kb-rag PDF, this page has a normal text layer and an embedded image.")
    p.insert_image(pymupdf.Rect(72,100,172,200), stream=tiny_png())
    b = d.tobytes(); d.close(); return b

def repeated_image_pdf(pages=4):
    logo = tiny_png(); d = pymupdf.open()
    for _ in range(pages):
        p = d.new_page(); p.insert_text((72,72), "Hello kb-rag PDF, this page has a normal text layer and a header logo.")
        p.insert_image(pymupdf.Rect(72,100,172,200), stream=logo)
    b = d.tobytes(); d.close(); return b

def scanned_pdf(pages=1):
    d = pymupdf.open()
    for _ in range(pages): d.new_page()
    b = d.tobytes(); d.close(); return b

def docx_bytes():
    doc = Document(); doc.add_heading("Title", level=1); doc.add_paragraph("Hello kb-rag DOCX")
    t = doc.add_table(rows=2, cols=2)
    t.rows[0].cells[0].text="Name"; t.rows[0].cells[1].text="Age"
    t.rows[1].cells[0].text="Alice"; t.rows[1].cells[1].text="30"
    buf = io.BytesIO(); doc.save(buf); return buf.getvalue()

def docx_with_image():
    doc = Document(); doc.add_paragraph("Hello kb-rag DOCX with an embedded image")
    doc.add_picture(io.BytesIO(tiny_png())); doc.add_paragraph("after the image")
    buf = io.BytesIO(); doc.save(buf); return buf.getvalue()

def xlsx_bytes():
    wb = Workbook(); s = wb.active; s.title="Sheet1"
    s.append(["Name","Age"]); s.append(["Alice",30]); s.append(["Bob",25])
    buf = io.BytesIO(); wb.save(buf); return buf.getvalue()

def csv_bytes():
    buf = io.StringIO(); w = csv.writer(buf)
    w.writerow(["Name","Age"]); w.writerow(["Alice","30"]); w.writerow(["Bob","25"])
    return buf.getvalue().encode()

def chat_csv(header, rows):
    buf = io.StringIO(); w = csv.writer(buf); w.writerow(header)
    for r in rows: w.writerow(r)
    return buf.getvalue().encode()

def chat_xlsx(header, rows):
    wb = Workbook(); s = wb.active; s.title="Sheet1"; s.append(list(header))
    for r in rows: s.append(list(r))
    buf = io.BytesIO(); wb.save(buf); return buf.getvalue()

SAMPLE_HTML = """<!DOCTYPE html><html><head><title>kb-rag 指南</title>
<style>body{color:red}</style><script>console.log("t");</script></head>
<body><h1>快速开始</h1><p>第一段，含 <a href="https://example.com/x">链接文本</a> 与 <b>加粗</b>。</p>
<h2>安装步骤</h2><ul><li>第一步</li><li>第二步</li></ul>
<script>alert("b");</script><noscript>请开启 JS</noscript></body></html>"""

CHAT_HEADER = ["room_name","NickName","Sender","IsSender","CreateTime","Type","StrContent"]

# ---------- 比对 ----------
results = []
def record(name, ok, detail=""):
    results.append((name, ok, detail))
    print(("  PASS  " if ok else "  DIFF  ") + name + (("  -> " + detail) if detail and not ok else ""))

def cmp_doc(name, filename, content, ext, compare_text=True):
    pj = parse_doc(PY_URL, filename, content, ext)
    jj = parse_doc(JAVA_URL, filename, content, ext)
    diffs = []
    if pj["code"] != jj["code"]: diffs.append(f"code {pj['code']} vs {jj['code']}")
    p, j = pj.get("data"), jj.get("data")
    if p is None or j is None:
        record(name, not diffs, "; ".join(diffs) or "both failed identically"); return
    if len(p["pages"]) != len(j["pages"]): diffs.append(f"pages {len(p['pages'])} vs {len(j['pages'])}")
    for a, b in zip(p["pages"], j["pages"]):
        if a["page_no"] != b["page_no"]: diffs.append(f"page_no {a['page_no']} vs {b['page_no']}")
        if a["scanned"] != b["scanned"]: diffs.append(f"p{a['page_no']} scanned {a['scanned']} vs {b['scanned']}")
        if a.get("ocr_source") != b.get("ocr_source"): diffs.append(f"p{a['page_no']} ocr_source differs")
        if compare_text and a["text"].split() != b["text"].split():
            diffs.append(f"p{a['page_no']} text: {a['text'][:60]!r} vs {b['text'][:60]!r}")
    if len(p["images"]) != len(j["images"]): diffs.append(f"images {len(p['images'])} vs {len(j['images'])}")
    for a, b in zip(p["images"], j["images"]):
        if (a["image_id"], a["page_no"], a["kind"]) != (b["image_id"], b["page_no"], b["kind"]):
            diffs.append(f"image meta {a['image_id']}/{a['kind']} vs {b['image_id']}/{b['kind']}")
    if p["markdown"].count("[[IMAGE:") != j["markdown"].count("[[IMAGE:"):
        diffs.append("placeholder count differs")
    # pages[].markdown 拼回 == markdown 的不变量，两边都必须成立
    for label, d in (("py", p), ("java", j)):
        if "\n\n".join(x["markdown"] for x in d["pages"]) != d["markdown"]:
            diffs.append(f"{label} page-markdown reassembly broken")
    if len(p["warnings"]) != len(j["warnings"]): diffs.append(f"warnings {len(p['warnings'])} vs {len(j['warnings'])}")
    record(name, not diffs, "; ".join(diffs))

def cmp_chat(name, filename, content, ext, profile=None, yml=None):
    pj = parse_chat(PY_URL, filename, content, ext, profile, yml)
    jj = parse_chat(JAVA_URL, filename, content, ext, profile, yml)
    diffs = []
    if pj["code"] != jj["code"]: diffs.append(f"code {pj['code']} vs {jj['code']}")
    p, j = pj.get("data"), jj.get("data")
    if p is None or j is None:
        record(name, not diffs, "; ".join(diffs) or "both failed identically"); return
    if p["skipped"] != j["skipped"]: diffs.append(f"skipped {p['skipped']} vs {j['skipped']}")
    if len(p["sessions"]) != len(j["sessions"]): diffs.append(f"sessions {len(p['sessions'])} vs {len(j['sessions'])}")
    for sa, sb in zip(p["sessions"], j["sessions"]):
        if (sa["session_id"], sa["session_name"]) != (sb["session_id"], sb["session_name"]):
            diffs.append(f"session id/name {sa['session_id']}/{sa['session_name']} vs {sb['session_id']}/{sb['session_name']}")
        if len(sa["messages"]) != len(sb["messages"]):
            diffs.append(f"messages {len(sa['messages'])} vs {len(sb['messages'])}"); continue
        for ma, mb in zip(sa["messages"], sb["messages"]):
            if ma != mb: diffs.append(f"message {ma} vs {mb}")
    record(name, not diffs, "; ".join(diffs))

print("=== /api/v1/parse ===")
cmp_doc("txt", "sample.txt", b"Hello kb-rag TXT\nsecond line", "txt")
cmp_doc("md", "sample.md", "# Heading\n\nHello kb-rag MD".encode(), "md")
cmp_doc("sql", "schema.sql", b"SELECT id, name FROM t_kb\nWHERE status = 'READY';", "sql")
cmp_doc("txt gbk", "gbk.txt", "中文编码测试".encode("gbk"), "txt")
cmp_doc("csv", "sample.csv", csv_bytes(), "csv")
cmp_doc("xlsx", "sample.xlsx", xlsx_bytes(), "xlsx")
cmp_doc("docx", "sample.docx", docx_bytes(), "docx")
cmp_doc("docx + image", "with_image.docx", docx_with_image(), "docx")
cmp_doc("html", "guide.html", SAMPLE_HTML.encode(), "html")
cmp_doc("htm", "legacy.htm", b"<p>htm works</p>", "htm")
cmp_doc("html malformed", "broken.html", b"<p>unclosed <b>bold</p></i><div>tail", "html")
# &nbsp; 是真实网页里最常见的隐形字符，而 Java 的 \s 与 String.strip() 都不认它——
# 两边对它的处理若不一致，语料里就会多出看不见的字符。
cmp_doc("html nbsp", "nbsp.html",
    "<html><body><h1>&nbsp;标题&nbsp;</h1><p>a&nbsp;&nbsp;b</p></body></html>".encode(), "html")
cmp_doc("pdf", "sample.pdf", pdf_bytes(), "pdf")
cmp_doc("pdf multi-page", "multi.pdf", multi_pdf(3), "pdf")
cmp_doc("pdf + embedded image", "with_image.pdf", pdf_with_image(), "pdf")
cmp_doc("pdf repeated logo (dedup)", "logo.pdf", repeated_image_pdf(4), "pdf")
cmp_doc("pdf scanned", "scanned.pdf", scanned_pdf(1), "pdf")
cmp_doc("pdf scanned x5", "scanned5.pdf", scanned_pdf(5), "pdf")
cmp_doc("unsupported ext", "x.exe", b"whatever", "exe")
cmp_doc("corrupt pdf", "broken.pdf", b"not a pdf at all", "pdf")

print("=== /api/v1/parse/chat ===")
cmp_chat("chat csv", "chat.csv", chat_csv(CHAT_HEADER, [
    ["room_a","Alice's Room","alice","1","1737800000","1","hello there"],
    ["room_a","Alice's Room","bob","0","1737800060","1","hi alice"]]), "csv")
cmp_chat("chat csv voice/video skip", "chat.csv", chat_csv(CHAT_HEADER, [
    ["room_a","Room","alice","1","1737800000","1","text message"],
    ["room_a","Room","alice","1","1737800001","34","[voice]"],
    ["room_a","Room","alice","1","1737800002","43","[video]"],
    ["room_a","Room","alice","1","1737800003","3","[image]"]]), "csv")
cmp_chat("chat csv bad time", "chat.csv", chat_csv(CHAT_HEADER, [
    ["room_a","Room","alice","1","not-a-time","1","bad"],
    ["room_a","Room","alice","1","1737800000","1","good"]]), "csv")
cmp_chat("chat csv multi-room", "chat.csv", chat_csv(CHAT_HEADER, [
    ["room_b","B","alice","1","1737800000","1","in b"],
    ["room_a","A","bob","0","1737800001","1","in a"],
    ["room_b","B","carol","0","1737800002","1","also in b"]]), "csv")
cmp_chat("chat csv missing content col", "chat.csv", chat_csv(CHAT_HEADER[:-1], [
    ["room_a","Room","alice","1","1737800000","1"]]), "csv")
cmp_chat("chat xlsx", "chat.xlsx", chat_xlsx(CHAT_HEADER, [
    ["room_b","Team Chat","carol","1","1737800000000","1","xlsx hello"]]), "xlsx")
cmp_chat("chat txt liuhen", "chat.txt",
    "2024-01-01 10:00:00 张三\n你好，最近怎么样？\n\n2024-01-01 10:05:00 李四\n挺好的，你呢？\n第二行内容\n".encode(), "txt")
cmp_chat("chat txt wechat_pc", "chat.txt",
    "张三 (2024-01-01 10:00:00): 你好\n李四 (2024-01-01 10:05:00): 挺好的，你呢？\n".encode(), "txt")
cmp_chat("chat txt wrong format", "notes.txt",
    "This is just some random text file.\nIt has multiple lines.\nNone look like a chat log.\nNot even close.\n".encode(), "txt")
cmp_chat("chat txt bad timestamp", "chat.txt",
    "2024-13-99 99:99:99 张三\nheader shape ok but date unreal\n\n2024-01-01 10:00:00 李四\nvalid\n".encode(), "txt")
cmp_chat("chat txt custom regex (profile_yaml)", "chat.txt",
    "[2024-01-01 10:00:00] alice >> hello there\n".encode(), "txt", None,
    "txt:\n  patterns:\n    - name: custom\n      regex: '^\\[(?P<send_time>\\d{4}-\\d{2}-\\d{2} \\d{2}:\\d{2}:\\d{2})\\] (?P<sender>\\S+) >> (?P<content>.*)$'\n")
HTML_CHAT = """<html><body>
<div class="message"><span class="sender">张三</span><span class="time">2024-01-01 10:00:00</span><div class="content">你好，最近怎么样？</div></div>
<div class="message"><span class="sender">李四</span><span class="time">2024-01-01 10:05:00</span><div class="content">挺好的，你呢？</div></div>
</body></html>"""
cmp_chat("chat html liuhen", "chat.html", HTML_CHAT.encode(), "html")
cmp_chat("chat html image placeholder", "chat.html",
    '<html><body><div class="message"><span class="sender">张三</span><span class="time">2024-01-01 10:00:00</span><div class="content"><img src="http://example.com/x.png"/></div></div></body></html>'.encode(), "html")
cmp_chat("chat html voice/video skip", "chat.html",
    '<html><body><div class="message"><span class="sender">A</span><span class="time">2024-01-01 10:00:00</span><div class="content">text</div></div><div class="message"><span class="sender">B</span><span class="time">2024-01-01 10:01:00</span><div class="content"><audio src="v.amr"></audio></div></div><div class="message"><span class="sender">C</span><span class="time">2024-01-01 10:02:00</span><div class="content"><video src="c.mp4"></video></div></div></body></html>'.encode(), "html")
cmp_chat("chat html script/style strip", "chat.html",
    '<html><body><script>alert(1)</script><style>.m{}</style><div class="message"><span class="sender">张三</span><span class="time">2024-01-01 10:00:00</span><div class="content">hello<script>evil()</script></div></div></body></html>'.encode(), "html")
cmp_chat("chat html custom selectors", "chat.html",
    '<html><body><li class="chat-row"><b class="who">alice</b><i class="at">2024-01-01 10:00:00</i><p class="msg">hi there</p></li></body></html>'.encode(), "html",
    None, "html:\n  message: li.chat-row\n  sender: b.who\n  time: i.at\n  content: p.msg\n")
cmp_chat("chat html nbsp padding", "chat.html",
    '<html><body><div class="message"><span class="sender">&nbsp;张三&nbsp;</span><span class="time">&nbsp;2024-01-01 10:00:00&nbsp;</span><div class="content">&nbsp;hello&nbsp;</div></div></body></html>'.encode(), "html")
cmp_chat("chat csv nbsp header", "chat.csv",
    chat_csv(["room\u00a0name", "create\u00a0time", "Str\u00a0Content"], [["room_a", "1737800000", "hi"]]), "csv",
    None, "session_id:\n  - roomname\nsend_time:\n  - CreateTime\ncontent:\n  - strcontent\n")
cmp_chat("chat html selector matches nothing", "chat.html",
    b"<html><body><div class='not-a-message'>x</div></body></html>", "html")
cmp_chat("chat profile_yaml over missing local", "chat.csv",
    chat_csv(["from","when","body"], [["alice","1737800000","hello via inline profile"]]), "csv",
    "this_profile_does_not_exist", "sender:\n  - from\nsend_time:\n  - when\ncontent:\n  - body\n")
cmp_chat("chat unknown profile", "chat.csv", chat_csv(CHAT_HEADER, [
    ["room_a","Room","alice","1","1737800000","1","hi"]]), "csv", "does_not_exist")
cmp_chat("chat unsupported ext", "chat.pdf", b"x", "pdf")

passed = sum(1 for _, ok, _ in results if ok)
print(f"\n=== {passed}/{len(results)} identical ===")
sys.exit(0 if passed == len(results) else 1)
