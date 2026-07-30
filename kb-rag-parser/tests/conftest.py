"""Shared pytest fixtures: a TestClient plus minimal sample-file builders.

Every sample is generated in code (no binary fixtures committed to the
repo) using the same libraries the parsers themselves rely on, per the M1
verification requirement (extended for M3 multimodal/chat parsing).
"""

import csv
import io
from typing import Optional, Sequence

import pymupdf
import pytest
from docx import Document
from fastapi.testclient import TestClient
from openpyxl import Workbook

from app.main import app


@pytest.fixture()
def client() -> TestClient:
    return TestClient(app)


# Well above SCANNED_PAGE_TEXT_THRESHOLD's default of 20, so samples using
# this text are never mistaken for a scanned page unless a test explicitly
# wants that (see make_scanned_pdf_bytes).
def make_pdf_bytes(text: str = "Hello kb-rag PDF, this page has a normal text layer.") -> bytes:
    """Build a minimal one-page PDF using pymupdf (no reportlab available)."""
    document = pymupdf.open()
    page = document.new_page()
    page.insert_text((72, 72), text)
    data = document.tobytes()
    document.close()
    return data


def make_tiny_png_bytes() -> bytes:
    """A tiny valid PNG, rendered via pymupdf so no extra imaging library
    (e.g. Pillow) is needed just to build test fixtures."""
    document = pymupdf.open()
    page = document.new_page(width=40, height=40)
    png_bytes = page.get_pixmap().tobytes("png")
    document.close()
    return png_bytes


def make_pdf_with_embedded_image_bytes(
    text: str = "Hello kb-rag PDF, this page has a normal text layer and an embedded image.",
) -> bytes:
    """A one-page PDF with a real text layer (well above the scanned-page
    threshold) plus one embedded raster image."""
    document = pymupdf.open()
    page = document.new_page()
    page.insert_text((72, 72), text)
    rect = pymupdf.Rect(72, 100, 172, 200)
    page.insert_image(rect, stream=make_tiny_png_bytes())
    data = document.tobytes()
    document.close()
    return data


def make_pdf_with_repeated_image_bytes(
    pages: int = 3,
    text: str = "Hello kb-rag PDF, this page has a normal text layer and a header logo.",
) -> bytes:
    """A multi-page PDF whose pages all draw the *same* raster, the way a
    header/footer logo appears throughout a real document.

    pymupdf stores an identical stream once and lets every page reference
    that one xref (verified: all pages report the same xref), which is
    exactly the shape the parser's per-xref deduplication keys on.
    """
    logo = make_tiny_png_bytes()
    document = pymupdf.open()
    for _ in range(pages):
        page = document.new_page()
        page.insert_text((72, 72), text)
        page.insert_image(pymupdf.Rect(72, 100, 172, 200), stream=logo)
    data = document.tobytes()
    document.close()
    return data


def make_scanned_pdf_bytes(pages: int = 1) -> bytes:
    """A PDF with no text layer at all, simulating scanned pages."""
    document = pymupdf.open()
    for _ in range(pages):
        document.new_page()
    data = document.tobytes()
    document.close()
    return data


def make_scanned_pdf_with_ocrable_text_bytes(text: str = "HELLO OCR") -> bytes:
    """A one-page PDF with zero native text layer (page.get_text() == "",
    so it is still detected as scanned=True) whose rendered pixmap
    nonetheless contains legible glyph pixels, for testing real PaddleOCR
    inference end to end (M8-CONTRACTS.md §0.4): a helper document renders
    `text` with a real (extractable) text layer, then its pixmap -- a pure
    raster image of those glyphs -- is embedded as a picture on the actual
    test page, which itself never gets a text layer of its own."""
    helper_document = pymupdf.open()
    helper_page = helper_document.new_page(width=300, height=100)
    helper_page.insert_text((10, 60), text, fontsize=36)
    png_bytes = helper_page.get_pixmap(dpi=200).tobytes("png")
    helper_document.close()

    document = pymupdf.open()
    page = document.new_page(width=300, height=100)
    page.insert_image(pymupdf.Rect(0, 0, 300, 100), stream=png_bytes)
    data = document.tobytes()
    document.close()
    return data


def make_docx_bytes(heading: str = "Title", paragraph: str = "Hello kb-rag DOCX") -> bytes:
    document = Document()
    document.add_heading(heading, level=1)
    document.add_paragraph(paragraph)
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Name"
    table.rows[0].cells[1].text = "Age"
    table.rows[1].cells[0].text = "Alice"
    table.rows[1].cells[1].text = "30"

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def make_docx_with_image_bytes(paragraph: str = "Hello kb-rag DOCX with an embedded image") -> bytes:
    document = Document()
    document.add_paragraph(paragraph)
    document.add_picture(io.BytesIO(make_tiny_png_bytes()))
    document.add_paragraph("after the image")

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def make_xlsx_bytes() -> bytes:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Sheet1"
    sheet.append(["Name", "Age"])
    sheet.append(["Alice", 30])
    sheet.append(["Bob", 25])

    buffer = io.BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


def make_csv_bytes() -> bytes:
    buffer = io.StringIO()
    writer = csv.writer(buffer)
    writer.writerow(["Name", "Age"])
    writer.writerow(["Alice", "30"])
    writer.writerow(["Bob", "25"])
    return buffer.getvalue().encode("utf-8")


def make_chat_csv_bytes(header: Sequence[str], rows: Sequence[Sequence[object]]) -> bytes:
    """Build a chat-log csv export with an arbitrary header/rows shape, so
    tests can exercise different column-naming and value conventions."""
    buffer = io.StringIO()
    writer = csv.writer(buffer)
    writer.writerow(header)
    for row in rows:
        writer.writerow(row)
    return buffer.getvalue().encode("utf-8")


def make_chat_xlsx_bytes(header: Sequence[str], rows: Sequence[Sequence[object]]) -> bytes:
    """Same as make_chat_csv_bytes but as an .xlsx workbook (single sheet)."""
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Sheet1"
    sheet.append(list(header))
    for row in rows:
        sheet.append(list(row))

    buffer = io.BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


def post_parse(client: TestClient, filename: str, content: bytes, file_ext: str):
    return client.post(
        "/api/v1/parse",
        files={"file": (filename, content, "application/octet-stream")},
        data={"file_ext": file_ext},
    )


def post_parse_chat(
    client: TestClient,
    filename: str,
    content: bytes,
    file_ext: str,
    mapping_profile: Optional[str] = None,
    profile_yaml: Optional[str] = None,
):
    form_data = {"file_ext": file_ext}
    if mapping_profile is not None:
        form_data["mapping_profile"] = mapping_profile
    if profile_yaml is not None:
        form_data["profile_yaml"] = profile_yaml
    return client.post(
        "/api/v1/parse/chat",
        files={"file": (filename, content, "application/octet-stream")},
        data=form_data,
    )
