"""DOCX parser backed by python-docx.

Security note: docx is a zip package containing XML parts. Before handing
the bytes to python-docx (or reading any zip entry directly, see the M3
image extraction below), ``ensure_zip_is_safe`` runs the zip-slip /
zip-bomb precheck (requirement §4.2). XXE hardening for the XML parts
themselves is already provided by python-docx's own lxml parser
configuration (``resolve_entities=False`` -- audited in app/security.py's
module docstring); the relationship XML this module parses directly (see
``_extract_docx_media``) goes through ``defusedxml.ElementTree`` explicitly
so it is safe regardless of whether ``app.security.harden_xml_parsing`` has
already run in the current process.

Author: owlzhangfq@gmail.com
"""

import io
import logging
import posixpath
import re
import zipfile
from typing import Dict, Iterator, List, Set, Tuple, Union

import defusedxml.ElementTree as safe_et
from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from app.errors import ErrorCode, ParseError
from app.models import PageContent, ParseData
from app.parsers.base import BaseParser
from app.parsers.images import KIND_EMBEDDED, ImageAssetCollector, guess_media_type, image_placeholder
from app.security import ensure_zip_is_safe

logger = logging.getLogger(__name__)

# Word style names that indicate a heading level; used to render markdown "#" prefixes.
_HEADING_STYLE_PREFIX = "Heading"

# docx has no reliable page-boundary information (see DocxParser docstring),
# so every embedded image is attributed to this single logical page.
_SINGLE_PAGE_NO = 1

# Where embedded pictures live inside the docx zip package, and the
# relationship part that maps a run's r:embed id to one of those files.
_MEDIA_PREFIX = "word/media/"
_DOCUMENT_RELS_PATH = "word/_rels/document.xml.rels"
_IMAGE_RELATIONSHIP_TYPE_SUFFIX = "/image"

# A run's inline/floating picture references its image via
# <a:blip r:embed="rIdN"/> inside <w:drawing>; a plain regex over the
# paragraph's serialized XML is far simpler than walking the OOXML drawing
# schema and is sufficient since we only need the referenced relationship
# ids, in document order.
_EMBED_ID_PATTERN = re.compile(r'r:embed="([^"]+)"')


def _iter_block_items(document: Document) -> Iterator[Union[Paragraph, Table]]:
    """Yield paragraphs and tables in the document's original reading order.

    python-docx exposes ``document.paragraphs`` and ``document.tables``
    separately with no ordering between them; walking the body XML directly
    is the standard recipe to reconstruct true document order.
    """
    body = document.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def _paragraph_to_markdown(paragraph: Paragraph) -> str:
    style_name = paragraph.style.name if paragraph.style is not None else ""
    text = paragraph.text
    if style_name.startswith(_HEADING_STYLE_PREFIX):
        level_str = style_name[len(_HEADING_STYLE_PREFIX) :].strip()
        level = int(level_str) if level_str.isdigit() else 1
        level = min(max(level, 1), 6)
        return f"{'#' * level} {text}"
    return text


def _table_to_markdown(table: Table) -> str:
    rows = [[cell.text for cell in row.cells] for row in table.rows]
    if not rows:
        return ""
    header, *body_rows = rows
    lines = ["| " + " | ".join(header) + " |", "| " + " | ".join(["---"] * len(header)) + " |"]
    for row in body_rows:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


def _ext_of(zip_path: str) -> str:
    return zip_path.rsplit(".", 1)[-1] if "." in zip_path else ""


def _extract_docx_media(content: bytes) -> Tuple[Dict[str, bytes], Dict[str, str]]:
    """Read every embedded media file plus the r:embed -> media-path mapping
    directly from the docx zip package.

    Returns (media_bytes_by_path, rid_to_media_path):
      - media_bytes_by_path: every ``word/media/*`` entry, keyed by its
        full zip path -- this is the exhaustive set of embedded images,
        independent of how/where they are anchored (inline, floating,
        header/footer).
      - rid_to_media_path: for relationships whose Type is an image
        relationship, maps the relationship id (as referenced by
        ``r:embed`` in a run) to the same zip path, so callers can place a
        placeholder at the paragraph where that image actually occurs.
    """
    media_bytes_by_path: Dict[str, bytes] = {}
    rid_to_media_path: Dict[str, str] = {}
    with zipfile.ZipFile(io.BytesIO(content)) as archive:
        names = archive.namelist()
        for name in names:
            if name.startswith(_MEDIA_PREFIX):
                media_bytes_by_path[name] = archive.read(name)

        if _DOCUMENT_RELS_PATH in names:
            rels_root = safe_et.fromstring(archive.read(_DOCUMENT_RELS_PATH))
            for rel in rels_root:
                rel_id = rel.get("Id")
                target = rel.get("Target")
                rel_type = rel.get("Type", "")
                if rel_id and target and rel_type.endswith(_IMAGE_RELATIONSHIP_TYPE_SUFFIX):
                    rid_to_media_path[rel_id] = posixpath.normpath(posixpath.join("word", target))
    return media_bytes_by_path, rid_to_media_path


def _collect_paragraph_image_placeholders(
    paragraph: Paragraph,
    rid_to_media_path: Dict[str, str],
    media_bytes_by_path: Dict[str, bytes],
    consumed_paths: Set[str],
    collector: ImageAssetCollector,
) -> List[str]:
    """Emit one placeholder line per image this paragraph references, in order."""
    placeholders: List[str] = []
    for rid in _EMBED_ID_PATTERN.findall(paragraph._p.xml):
        media_path = rid_to_media_path.get(rid)
        if media_path is None or media_path in consumed_paths:
            continue
        raw_bytes = media_bytes_by_path.get(media_path)
        if raw_bytes is None:
            continue
        consumed_paths.add(media_path)
        image_id = collector.try_add(
            _SINGLE_PAGE_NO, kind=KIND_EMBEDDED, media_type=guess_media_type(_ext_of(media_path)), raw_bytes=raw_bytes
        )
        if image_id:
            placeholders.append(image_placeholder(image_id))
    return placeholders


def _collect_leftover_image_placeholders(
    media_bytes_by_path: Dict[str, bytes],
    consumed_paths: Set[str],
    collector: ImageAssetCollector,
) -> List[str]:
    """Placeholders for media files no paragraph scan matched (e.g. images
    anchored in a table cell, header, or footer): appended once at the end
    so no embedded image is silently dropped from images[]."""
    placeholders: List[str] = []
    for media_path in sorted(media_bytes_by_path):
        if media_path in consumed_paths:
            continue
        raw_bytes = media_bytes_by_path[media_path]
        image_id = collector.try_add(
            _SINGLE_PAGE_NO, kind=KIND_EMBEDDED, media_type=guess_media_type(_ext_of(media_path)), raw_bytes=raw_bytes
        )
        if image_id:
            placeholders.append(image_placeholder(image_id))
    return placeholders


class DocxParser(BaseParser):
    """Extracts full-document text (paragraphs + tables) and embedded
    images from a .docx file.

    Word documents have no reliable page-boundary information without a
    layout engine (page breaks depend on rendering, fonts, and margins), so
    this returns the whole document as a single logical page (page_no=1);
    every extracted image is likewise attributed to page_no=1
    (M3-CONTRACTS.md §2.1). Real pagination is a TODO if a downstream
    consumer needs it.
    """

    def parse(self, content: bytes, filename: str) -> ParseData:
        ensure_zip_is_safe(content)

        try:
            document = Document(io.BytesIO(content))
        except Exception as exc:
            logger.error(
                "docx parse failed, errorCode=%s, filename=%s, stage=open",
                ErrorCode.PARSE_FAILED,
                filename,
            )
            raise ParseError(f"failed to open docx: {exc}") from exc

        try:
            media_bytes_by_path, rid_to_media_path = _extract_docx_media(content)
            consumed_paths: Set[str] = set()
            collector = ImageAssetCollector()

            blocks = list(_iter_block_items(document))
            markdown_parts: List[str] = []
            plain_text_parts: List[str] = []
            for block in blocks:
                if isinstance(block, Paragraph):
                    markdown_parts.append(_paragraph_to_markdown(block))
                    plain_text_parts.append(block.text)
                    markdown_parts.extend(
                        _collect_paragraph_image_placeholders(
                            block, rid_to_media_path, media_bytes_by_path, consumed_paths, collector
                        )
                    )
                elif isinstance(block, Table):
                    markdown_parts.append(_table_to_markdown(block))

            markdown_parts.extend(
                _collect_leftover_image_placeholders(media_bytes_by_path, consumed_paths, collector)
            )

            markdown = "\n\n".join(part for part in markdown_parts if part)
            pages = [
                PageContent(
                    page_no=_SINGLE_PAGE_NO,
                    text="\n".join(plain_text_parts),
                    markdown=markdown,
                    scanned=False,
                )
            ]
            return ParseData(
                markdown=markdown, pages=pages, images=collector.images, warnings=collector.warnings
            )
        except Exception as exc:
            logger.error(
                "docx parse failed, errorCode=%s, filename=%s, stage=extract",
                ErrorCode.PARSE_FAILED,
                filename,
            )
            raise ParseError(f"failed to extract docx content: {exc}") from exc
