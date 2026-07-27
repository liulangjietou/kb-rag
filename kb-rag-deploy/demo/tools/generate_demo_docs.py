"""Generate the kb-rag demo document set (docx / pdf / xlsx) shipped under demo/docs/.

This script is the reproducible source for three of the four Demo documents referenced
by demo/manifest.json (the fourth, the Markdown file, is plain text and hand-authored,
no generation needed). All document content below is original text written for this
project describing RAG / knowledge-base concepts; none of it is copied from any
third-party or copyrighted source, and none of it contains real personal data
(demo/eval-cases.json quotes verbatim spans from this content, so keep the Chinese
strings below stable across regenerations unless you also update eval-cases.json).

Usage:
    pip install -r requirements.txt
    python3 generate_demo_docs.py

Outputs (relative to this file's parent's parent, i.e. demo/):
    docs/02-document-parsing-and-chunking.docx
    docs/03-hybrid-retrieval-and-rerank.pdf   (embeds one self-drawn PNG diagram,
                                                used to exercise the VLM image pipeline
                                                — see M3-CONTRACTS.md §2.1/§3.2)
    docs/04-eval-metrics-comparison.xlsx

Reproducibility note: office/PDF containers normally embed a "created/modified"
timestamp that changes on every run even when the visible content is byte-identical,
which would silently break demo/eval-cases.json's content_hash_sha256 (computed over
the whole file) on every regeneration. This script pins those timestamps to a fixed
constant (FIXED_TIMESTAMP below) and disables reportlab's random PDF file ID
(`invariant=1`) so re-running with unchanged text/table content reproduces
byte-identical output. If you DO change the visible content, the hash still changes
as expected — update demo/eval-cases.json's `content_hash_sha256` (and any `span`
values that no longer match verbatim) in the same change.

Author: owlzhangfq@gmail.com
"""

from __future__ import annotations

import datetime
import os
import re
import tempfile
import zipfile

from docx import Document
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
from PIL import Image as PILImage
from PIL import ImageDraw, ImageFont
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import Image as RLImage
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

DEMO_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DOCS_DIR = os.path.join(DEMO_DIR, "docs")

# Fixed creation/modification timestamp embedded into docx/xlsx/pdf metadata so that
# regenerating with unchanged content yields byte-identical files (see module docstring).
FIXED_TIMESTAMP = datetime.datetime(2026, 7, 26, 0, 0, 0)
DOC_AUTHOR = "kb-rag-deploy demo tools"

# Candidate CJK-capable TrueType fonts to try, in order, for the PNG diagram
# (macOS system fonts first, then common Linux/Windows locations). If none is
# found the diagram falls back to ASCII-only labels so the script still runs
# on a bare CI image without failing.
_CJK_FONT_CANDIDATES = [
    "/System/Library/Fonts/STHeiti Medium.ttc",
    "/System/Library/Fonts/Supplemental/Songti.ttc",
    "/usr/share/fonts/truetype/noto/NotoSansCJKsc-Regular.otf",
    "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
    "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",
    "C:\\Windows\\Fonts\\msyh.ttc",
]


_ZIP_FIXED_DATE_TIME = (2026, 7, 26, 0, 0, 0)


def _normalize_zip_timestamps(path: str, content_patch=None) -> None:
    """Rewrite every per-entry timestamp inside a docx/xlsx (both OOXML zip containers)
    to a fixed value, optionally patching one entry's bytes first.

    python-docx / openpyxl only let you set the *document-level* core properties
    (author/created/modified in docProps/core.xml); the underlying `zipfile` module
    still stamps each individual zip entry (word/document.xml, xl/worksheets/sheet1.xml,
    ...) with the current local time, which makes two otherwise-identical runs produce
    different bytes (and therefore a different SHA-256). Rebuilding the zip with a fixed
    `date_time` on every entry (same content, same order, same compression) removes that
    last source of nondeterminism so demo/eval-cases.json's content_hash_sha256 stays
    valid across regenerations that don't change the visible content.

    `content_patch` is an optional `{entry_filename: (bytes) -> bytes}` mapping applied
    before rewriting — needed for xlsx because `openpyxl.writer.excel.save_workbook`
    unconditionally re-stamps `docProps/core.xml`'s `<dcterms:modified>` to
    `datetime.now()` at save time, ignoring whatever was set on `wb.properties.modified`
    beforehand (see openpyxl/writer/excel.py `save_workbook`).
    """
    content_patch = content_patch or {}
    with zipfile.ZipFile(path, "r") as src:
        entries = [(info, src.read(info.filename)) for info in src.infolist()]

    tmp_path = path + ".tmp"
    with zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as dst:
        for info, data in entries:
            if info.filename in content_patch:
                data = content_patch[info.filename](data)
            new_info = zipfile.ZipInfo(info.filename, date_time=_ZIP_FIXED_DATE_TIME)
            new_info.compress_type = zipfile.ZIP_DEFLATED
            new_info.external_attr = info.external_attr
            dst.writestr(new_info, data)
    os.replace(tmp_path, path)


def _find_cjk_font() -> str | None:
    for path in _CJK_FONT_CANDIDATES:
        if os.path.isfile(path):
            return path
    return None


def _load_font(path: str | None, size: int) -> ImageFont.FreeTypeFont:
    if path is None:
        return ImageFont.load_default()
    try:
        return ImageFont.truetype(path, size)
    except OSError:
        return ImageFont.load_default()


def _draw_box(draw: ImageDraw.ImageDraw, box, lines, font, fill, text_fill="#1f2933"):
    x0, y0, x1, y1 = box
    draw.rounded_rectangle(box, radius=14, fill=fill, outline="#334155", width=2)
    total_h = sum(draw.textbbox((0, 0), line, font=font)[3] for line in lines)
    line_gap = 6
    total_h += line_gap * (len(lines) - 1)
    cur_y = (y0 + y1) / 2 - total_h / 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        w = bbox[2] - bbox[0]
        h = bbox[3] - bbox[1]
        draw.text(((x0 + x1) / 2 - w / 2, cur_y), line, font=font, fill=text_fill)
        cur_y += h + line_gap


def _draw_h_arrow(draw: ImageDraw.ImageDraw, x0, x1, y, color="#475569"):
    draw.line((x0, y, x1, y), fill=color, width=4)
    draw.polygon([(x1, y - 8), (x1, y + 8), (x1 + 12, y)], fill=color)


def _draw_v_arrow(draw: ImageDraw.ImageDraw, x, y0, y1, color="#475569"):
    draw.line((x, y0, x, y1), fill=color, width=4)
    draw.polygon([(x - 8, y1), (x + 8, y1), (x, y1 + 12)], fill=color)


def generate_pipeline_diagram(output_path: str, cjk_font_path: str | None) -> None:
    """Draw an original two-row (snake-layout) pipeline diagram as a PNG.

    Content: upload -> parse -> clean -> chunk -> embed (row 1, left to right),
    then down into row 2: dual-write index -> hybrid recall -> rerank -> generate
    (row 2, right to left), matching the index/retrieval pipeline described in
    demo/docs/03-hybrid-retrieval-and-rerank.pdf. This is a hand-drawn diagram
    made for this Demo, not a screenshot of any real system or third-party asset.
    """
    width, height = 1400, 760
    img = PILImage.new("RGB", (width, height), "#f8fafc")
    draw = ImageDraw.Draw(img)

    title_font = _load_font(cjk_font_path, 34)
    box_font = _load_font(cjk_font_path, 24)
    caption_font = _load_font(cjk_font_path, 20)

    title = "RAG 索引与检索流水线示意图" if cjk_font_path else "RAG Index & Retrieval Pipeline"
    draw.text((40, 30), title, font=title_font, fill="#0f172a")

    box_w, box_h = 220, 130
    row1_y = 160
    row2_y = 460
    gap = 40
    start_x = 40

    row1_labels = [
        ["文档上传", "Upload"],
        ["解析", "Parse"],
        ["清洗", "Clean"],
        ["切分", "Chunk"],
        ["嵌入", "Embed"],
    ]
    row2_labels = [
        ["双写索引", "Milvus / ES"],
        ["混合检索", "Hybrid Recall"],
        ["重排序", "Rerank"],
        ["生成回答", "Generate"],
    ]
    if not cjk_font_path:
        row1_labels = [[b] for _, b in row1_labels]
        row2_labels = [[b] for _, b in row2_labels]

    row1_boxes = []
    x = start_x
    for labels in row1_labels:
        box = (x, row1_y, x + box_w, row1_y + box_h)
        row1_boxes.append(box)
        _draw_box(draw, box, labels, box_font, fill="#dbeafe")
        x += box_w + gap

    for i in range(len(row1_boxes) - 1):
        y_mid = row1_y + box_h / 2
        _draw_h_arrow(draw, row1_boxes[i][2], row1_boxes[i + 1][0], y_mid)

    row2_boxes = []
    x = row1_boxes[-1][2] - box_w
    for labels in row2_labels:
        box = (x, row2_y, x + box_w, row2_y + box_h)
        row2_boxes.append(box)
        _draw_box(draw, box, labels, box_font, fill="#dcfce7")
        x -= box_w + gap

    last_row1 = row1_boxes[-1]
    first_row2 = row2_boxes[0]
    mid_x = (last_row1[0] + last_row1[2]) / 2
    _draw_v_arrow(draw, mid_x, last_row1[3], first_row2[1])

    for i in range(len(row2_boxes) - 1):
        y_mid = row2_y + box_h / 2
        _draw_h_arrow(draw, row2_boxes[i][0], row2_boxes[i + 1][2], y_mid)

    caption = (
        "自造示意图，用于 kb-rag Demo 素材验证图片解析链路，非真实系统截图"
        if cjk_font_path
        else "Original diagram for kb-rag demo assets, not a real system screenshot"
    )
    draw.text((40, height - 40), caption, font=caption_font, fill="#475569")

    img.save(output_path, format="PNG")


def generate_docx(output_path: str) -> None:
    doc = Document()
    doc.add_heading("文档解析与分片切分策略实践", level=1)

    doc.add_heading("多格式解析与版面判定", level=2)
    doc.add_paragraph(
        "知识库要处理的原始文件五花八门：可直接抽取文字的电子版 PDF、Word、Excel、纯文本，"
        "也有扫描件、截图、聊天记录导出这类没有文本层或结构松散的资料。解析服务的第一个职责"
        "是识别版面：对能直接抽取文本的页面，逐段提取正文与表格；对整页没有文本层的扫描页，"
        "则把该页整体渲染成图片，连同占位符一起交给下游做图文理解，文字内容由视觉大模型转录"
        "得到，而不是本地维护一套体积庞大的离线 OCR 模型。"
    )

    doc.add_heading("图片与占位符的对应关系", level=2)
    doc.add_paragraph(
        "无论是文档内嵌图片还是扫描页渲染出的图片，解析阶段都只产出文本、占位符与图片二进制，"
        "并不调用任何大模型。占位符在正文中独占一行，下游服务按图片 ID 精确定位插入点，把图片"
        "的文字描述和转录文本换回占位符所在位置，让图片语义自然地参与到后续统一切分中，而不是"
        "把图片孤立成一个缺少上下文的片段。"
    )

    doc.add_heading("切分策略的选择", level=2)
    doc.add_paragraph(
        "拿到清洗后的文本，还需要选一种合适的切分策略：结构清晰的长文档适合按标题层级切分，"
        "保留父子关系与层级信息；表格型数据按行切分并把表头拼接进每一行，避免脱离表头后数字"
        "失去含义；无结构的长文本（例如聊天记录聚合后的大段内容）更适合语义切分或大模型判断"
        "切分点。切分策略与父子分片是两个正交的开关，二者可以自由组合，用评测结果说话，而不是"
        "凭经验一次定死。"
    )

    for paragraph in doc.paragraphs:
        if paragraph.style.name == "Normal":
            for run in paragraph.runs:
                run.font.size = Pt(11)

    core_props = doc.core_properties
    core_props.author = DOC_AUTHOR
    core_props.last_modified_by = DOC_AUTHOR
    core_props.created = FIXED_TIMESTAMP
    core_props.modified = FIXED_TIMESTAMP
    core_props.revision = 1

    doc.save(output_path)
    _normalize_zip_timestamps(output_path)


def generate_pdf(output_path: str) -> None:
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))

    styles = {
        "title": ParagraphStyle(
            "title", fontName="STSong-Light", fontSize=18, leading=24, spaceAfter=14
        ),
        "h2": ParagraphStyle(
            "h2", fontName="STSong-Light", fontSize=14, leading=20, spaceBefore=10, spaceAfter=6
        ),
        "body": ParagraphStyle(
            "body", fontName="STSong-Light", fontSize=10.5, leading=17, spaceAfter=8
        ),
        "caption": ParagraphStyle(
            "caption",
            fontName="STSong-Light",
            fontSize=9,
            leading=13,
            textColor="#475569",
            spaceBefore=6,
        ),
    }

    cjk_font_path = _find_cjk_font()
    with tempfile.TemporaryDirectory() as tmp_dir:
        diagram_path = os.path.join(tmp_dir, "pipeline-diagram.png")
        generate_pipeline_diagram(diagram_path, cjk_font_path)

        story = [
            Paragraph("混合检索与重排序技术白皮书", styles["title"]),
            Paragraph("双路召回：向量与全文各有所长", styles["h2"]),
            Paragraph(
                "向量检索把文本映射到高维语义空间，擅长捕捉意思相近但用词不同的相关性，例如"
                "多少钱和价格是多少在语义上很接近；全文检索基于关键词匹配和词频统计，擅长专有"
                "名词、型号、编号这类必须精确命中的场景。单独依赖任何一路都会有明显短板，因此"
                "生产级检索系统通常同时跑向量与全文两路召回，再把结果融合，取长补短。",
                styles["body"],
            ),
            Paragraph("融合与重排序", styles["h2"]),
            Paragraph(
                "两路候选集的原始分数量纲完全不同：余弦相似度落在有界区间，而全文检索的原始"
                "分数没有上界，直接相加会让全文分数淹没向量分数。常见做法是先对每路候选集做"
                "归一化，再按权重加权，或者使用倒数排名融合，不依赖具体分值，只看排名先后，"
                "天然免疫量纲问题。融合之后的候选集通常还会经过一个专门的重排模型二次打分，"
                "重排模型能同时看到问题和候选片段的完整语义交互，排序质量明显优于任何单一召回"
                "路径，但计算成本更高，因此只对经过初筛的候选集起作用，并设置合理的超时与降级"
                "策略，避免拖慢整体响应。",
                styles["body"],
            ),
            Paragraph("阈值与父子分片的返回粒度", styles["h2"]),
            Paragraph(
                "业务上还需要一个宁缺毋滥的开关：把最终得分低于阈值的结果过滤掉，避免返回明显"
                "不相关的内容。如果知识库启用了父子分片，检索用小片，展示回大片，阈值与最终"
                "返回条数就要作用在合并后的父片维度上，而不是子片维度，这样用户看到的答案上下"
                "文才是连贯完整的，不会出现同一段父片内容被多次重复计数的情况。下图展示了从"
                "原始文档到最终生成答案的完整链路。",
                styles["body"],
            ),
            Spacer(1, 6),
            RLImage(diagram_path, width=16 * cm, height=16 * cm * 760 / 1400),
            Paragraph(
                "图：文档从上传到生成答案的索引与检索流水线（自绘示意图，用于验证 VLM 图片理解链路）",
                styles["caption"],
            ),
        ]

        doc = SimpleDocTemplate(
            output_path,
            pagesize=A4,
            leftMargin=2 * cm,
            rightMargin=2 * cm,
            topMargin=2 * cm,
            bottomMargin=2 * cm,
            title="混合检索与重排序技术白皮书",
            author=DOC_AUTHOR,
            creator=DOC_AUTHOR,
            producer="reportlab",
            # invariant=1 fixes creation/mod dates and the internal file ID instead of
            # randomizing them per run, so unchanged content reproduces identical bytes.
            invariant=1,
        )
        doc.build(story)


def generate_xlsx(output_path: str) -> None:
    wb = Workbook()
    wb.properties.creator = DOC_AUTHOR
    wb.properties.lastModifiedBy = DOC_AUTHOR
    wb.properties.created = FIXED_TIMESTAMP
    wb.properties.modified = FIXED_TIMESTAMP

    note_sheet = wb.active
    note_sheet.title = "说明"
    note_sheet.sheet_view.showGridLines = False
    note_sheet.column_dimensions["A"].width = 100
    intro_lines = [
        "评测指标说明",
        "",
        "检索质量不能只凭感觉判断，需要一套可复现的量化指标。Recall@K 衡量前 K 条结果里覆盖了"
        "多少标注为相关的证据，是最核心的召回类指标；Precision@K 衡量前 K 条结果里有多少是真正"
        "相关的，反映结果的干净程度；Hit Rate 只关心前 K 条里是否至少命中一条相关证据，是一个"
        "更宽松的及格线指标；MRR 关注第一条相关结果出现的位置，位置越靠前分数越高，适合评估一次"
        "找到答案的场景；NDCG@K 在此基础上进一步考虑结果的排序质量，把命中的位置和相关程度都"
        "纳入打分。下方数据表对比了 BM25 单路、向量单路、混合检索、混合检索加重排序四种配置在"
        "同一评测集上的表现，数值仅为示例数据，用于演示评测报告的展示形式，并非真实压测结果。",
    ]
    for row_idx, line in enumerate(intro_lines, start=1):
        cell = note_sheet.cell(row=row_idx, column=1, value=line)
        cell.alignment = Alignment(wrap_text=True, vertical="top")
        if row_idx == 1:
            cell.font = Font(bold=True, size=14)
        else:
            cell.font = Font(size=11)
    note_sheet.row_dimensions[3].height = 160

    data_sheet = wb.create_sheet("评测结果对比")
    headers = ["检索配置", "Recall@5", "Precision@5", "Hit Rate", "MRR", "NDCG@5"]
    rows = [
        ["BM25 单路", 0.62, 0.41, 0.70, 0.55, 0.59],
        ["向量单路", 0.68, 0.44, 0.75, 0.60, 0.64],
        ["混合检索（RRF）", 0.79, 0.51, 0.86, 0.71, 0.75],
        ["混合检索 + 重排序", 0.88, 0.58, 0.93, 0.81, 0.85],
    ]
    header_fill = PatternFill(start_color="1e293b", end_color="1e293b", fill_type="solid")
    for col_idx, header in enumerate(headers, start=1):
        cell = data_sheet.cell(row=1, column=col_idx, value=header)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center")
    for row_idx, row in enumerate(rows, start=2):
        for col_idx, value in enumerate(row, start=1):
            data_sheet.cell(row=row_idx, column=col_idx, value=value)
    for col_idx in range(1, len(headers) + 1):
        data_sheet.column_dimensions[get_column_letter(col_idx)].width = 18

    wb.save(output_path)

    def _pin_core_modified(xml_bytes: bytes) -> bytes:
        fixed = FIXED_TIMESTAMP.strftime("%Y-%m-%dT%H:%M:%SZ").encode("ascii")
        return re.sub(
            rb"(<dcterms:modified[^>]*>)[^<]*(</dcterms:modified>)",
            lambda m: m.group(1) + fixed + m.group(2),
            xml_bytes,
        )

    _normalize_zip_timestamps(output_path, content_patch={"docProps/core.xml": _pin_core_modified})


def main() -> None:
    os.makedirs(DOCS_DIR, exist_ok=True)

    docx_path = os.path.join(DOCS_DIR, "02-document-parsing-and-chunking.docx")
    pdf_path = os.path.join(DOCS_DIR, "03-hybrid-retrieval-and-rerank.pdf")
    xlsx_path = os.path.join(DOCS_DIR, "04-eval-metrics-comparison.xlsx")

    print(f"[INFO] generating docx -> {docx_path}")
    generate_docx(docx_path)

    print(f"[INFO] generating pdf -> {pdf_path}")
    generate_pdf(pdf_path)

    print(f"[INFO] generating xlsx -> {xlsx_path}")
    generate_xlsx(xlsx_path)

    print("[INFO] demo document generation done")


if __name__ == "__main__":
    main()
