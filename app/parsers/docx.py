"""DOCX parser backed by python-docx.

Security note: docx is a zip package containing XML parts. Before handing
the bytes to python-docx, ``ensure_zip_is_safe`` runs the zip-slip /
zip-bomb precheck (requirement §4.2). XXE hardening for the XML parts
themselves is already provided by python-docx's own lxml parser
configuration (``resolve_entities=False`` — audited in app/security.py's
module docstring), so no extra patch is needed here.

Author: owlzhangfq@gmail.com
"""

import io
import logging
from typing import Iterator, Union

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from app.errors import ErrorCode, ParseError
from app.models import PageContent, ParseData
from app.parsers.base import BaseParser
from app.security import ensure_zip_is_safe

logger = logging.getLogger(__name__)

# Word style names that indicate a heading level; used to render markdown "#" prefixes.
_HEADING_STYLE_PREFIX = "Heading"


def _iter_block_items(document: Document) -> Iterator[Union[Paragraph, Table]]:
    """Yield paragraphs and tables in the document's original reading order.

    python-docx exposes ``document.paragraphs`` and ``document.tables``
    separately with no ordering between them; walking the body XML directly
    is the standard recipe to reconstruct true document order.
    """
    body = document.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def _paragraph_to_markdown(paragraph: Paragraph) -> str:
    style_name = paragraph.style.name if paragraph.style is not None else ""
    text = paragraph.text
    if style_name.startswith(_HEADING_STYLE_PREFIX):
        level_str = style_name[len(_HEADING_STYLE_PREFIX) :].strip()
        level = int(level_str) if level_str.isdigit() else 1
        level = min(max(level, 1), 6)
        return f"{'#' * level} {text}"
    return text


def _table_to_markdown(table: Table) -> str:
    rows = [[cell.text for cell in row.cells] for row in table.rows]
    if not rows:
        return ""
    header, *body_rows = rows
    lines = ["| " + " | ".join(header) + " |", "| " + " | ".join(["---"] * len(header)) + " |"]
    for row in body_rows:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


class DocxParser(BaseParser):
    """Extracts full-document text (paragraphs + tables) from a .docx file.

    Word documents have no reliable page-boundary information without a
    layout engine (page breaks depend on rendering, fonts, and margins), so
    M1 returns the whole document as a single logical page (page_no=1).
    Real pagination is a TODO if a downstream consumer needs it.
    """

    def parse(self, content: bytes, filename: str) -> ParseData:
        ensure_zip_is_safe(content)

        try:
            document = Document(io.BytesIO(content))
        except Exception as exc:
            logger.error(
                "docx parse failed, errorCode=%s, filename=%s, stage=open",
                ErrorCode.PARSE_FAILED,
                filename,
            )
            raise ParseError(f"failed to open docx: {exc}") from exc

        try:
            blocks = list(_iter_block_items(document))
            markdown_parts = []
            plain_text_parts = []
            for block in blocks:
                if isinstance(block, Paragraph):
                    markdown_parts.append(_paragraph_to_markdown(block))
                    plain_text_parts.append(block.text)
                elif isinstance(block, Table):
                    markdown_parts.append(_table_to_markdown(block))
            markdown = "\n\n".join(part for part in markdown_parts if part)
            pages = [PageContent(page_no=1, text="\n".join(plain_text_parts))]
            return ParseData(markdown=markdown, pages=pages, images=[])
        except Exception as exc:
            logger.error(
                "docx parse failed, errorCode=%s, filename=%s, stage=extract",
                ErrorCode.PARSE_FAILED,
                filename,
            )
            raise ParseError(f"failed to extract docx content: {exc}") from exc
